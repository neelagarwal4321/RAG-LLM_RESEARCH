from __future__ import annotations

import io
import textwrap
from pathlib import Path
from typing import Iterable, List

import fitz  # type: ignore
import pandas as pd
from docx import Document as DocxDocument  # type: ignore

from .types import DocumentSection, LoadedDocument


def load_documents(dataset: str, paths: Iterable[Path]) -> List[LoadedDocument]:
    documents: List[LoadedDocument] = []
    for path in paths:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            docs = _load_pdf(dataset, path)
        elif suffix == ".docx":
            docs = _load_docx(dataset, path)
        elif suffix in {".xlsx", ".xls", ".csv"}:
            docs = _load_tabular(dataset, path)
        elif suffix in {".txt", ".md"}:
            docs = _load_text(dataset, path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
        documents.extend(docs)
    return documents


def _load_pdf(dataset: str, path: Path) -> List[LoadedDocument]:
    sections: List[DocumentSection] = []
    doc = fitz.open(path)
    try:
        for page_index in range(len(doc)):
            page = doc[page_index]
            text = page.get_text("text")
            sections.append(
                DocumentSection(
                    text=text.strip(),
                    metadata={
                        "type": "pdf",
                        "page": page_index + 1,
                        "source_path": str(path),
                    },
                )
            )
    finally:
        doc.close()
    return [
        LoadedDocument(
            path=path,
            dataset=dataset,
            doc_id=_doc_id(dataset, path),
            sections=sections,
            metadata={"content_type": "pdf", "source_path": str(path)},
        )
    ]


def _load_docx(dataset: str, path: Path) -> List[LoadedDocument]:
    document = DocxDocument(str(path))
    sections: List[DocumentSection] = []
    for idx, para in enumerate(document.paragraphs):
        text = para.text.strip()
        if text:
            sections.append(
                DocumentSection(
                    text=text,
                    metadata={
                        "type": "docx",
                        "paragraph": idx + 1,
                        "source_path": str(path),
                    },
                )
            )
    return [
        LoadedDocument(
            path=path,
            dataset=dataset,
            doc_id=_doc_id(dataset, path),
            sections=sections,
            metadata={"content_type": "docx", "source_path": str(path)},
        )
    ]


def _load_tabular(dataset: str, path: Path) -> List[LoadedDocument]:
    sections: List[DocumentSection] = []
    suffix = path.suffix.lower()
    if suffix == ".csv":
        df = pd.read_csv(path)
        sections.extend(_tabular_sections(df, "CSV"))
    else:
        xls = pd.ExcelFile(path)
        for sheet_name in xls.sheet_names:
            df = xls.parse(sheet_name)
            sections.extend(_tabular_sections(df, sheet_name))
    return [
        LoadedDocument(
            path=path,
            dataset=dataset,
            doc_id=_doc_id(dataset, path),
            sections=sections,
            metadata={"content_type": "tabular", "source_path": str(path)},
        )
    ]


def _tabular_sections(df: pd.DataFrame, sheet_name: str) -> List[DocumentSection]:
    buffer = io.StringIO()
    df.to_csv(buffer, index=False)
    csv_text = buffer.getvalue()
    return [
        DocumentSection(
            text=csv_text.strip(),
            metadata={
                "type": "table",
                "sheet": sheet_name,
            },
        )
    ]


def _load_text(dataset: str, path: Path) -> List[LoadedDocument]:
    text = path.read_text(encoding="utf-8")
    sections = [
        DocumentSection(
            text=textwrap.dedent(text).strip(),
            metadata={
                "type": "text",
                "source_path": str(path),
            },
        )
    ]
    return [
        LoadedDocument(
            path=path,
            dataset=dataset,
            doc_id=_doc_id(dataset, path),
            sections=sections,
            metadata={"content_type": "text", "source_path": str(path)},
        )
    ]


def _doc_id(dataset: str, path: Path) -> str:
    return f"{dataset}:{path.stem}"
